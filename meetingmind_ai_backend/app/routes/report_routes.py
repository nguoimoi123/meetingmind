from flask import Blueprint, request, jsonify, send_file
from io import BytesIO
from docx import Document

report_bp = Blueprint("report", __name__, url_prefix="/report")


@report_bp.route("/docx", methods=["POST"])
def export_docx():
    data = request.get_json() or {}
    title = data.get("title") or "Meeting Report"
    summary = data.get("summary") or ""
    action_items = data.get("action_items") or []
    key_decisions = data.get("key_decisions") or []
    full_transcript = data.get("full_transcript") or ""

    if not any([summary.strip(), action_items, key_decisions, full_transcript.strip()]):
        return jsonify({"error": "No content to export"}), 400

    doc = Document()
    doc.add_heading(title, level=1)

    if summary.strip():
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(summary)

    if action_items:
        doc.add_heading("Action Items", level=2)
        for item in action_items:
            doc.add_paragraph(str(item), style="List Bullet")

    if key_decisions:
        doc.add_heading("Key Decisions", level=2)
        for item in key_decisions:
            doc.add_paragraph(str(item), style="List Bullet")

    if full_transcript.strip():
        doc.add_heading("Full Transcript", level=2)
        doc.add_paragraph(full_transcript)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="meeting_report.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
